"""Executable evidence for durable, sandboxed document analysis."""

from __future__ import annotations

import asyncio
import hashlib
import io
import json
import subprocess
import sys
import uuid
from collections.abc import AsyncIterator
from decimal import Decimal
from pathlib import Path
from typing import Any

import httpx
import pytest
import pytest_asyncio
from akc_api.main import create_app
from akc_api.models import (
    AnalysisTask,
    ArchitecturePlan,
    Block,
    BlueprintModule,
    Collection,
    CollectionFile,
    CollectionProcessingTaskBinding,
    CollectionSourceRoot,
    CreditLedger,
    Document,
    FeatureFlag,
    GpuProviderInvocation,
    ModelRegistry,
    OutboxEvent,
    Page,
    PageAsset,
    ProcessingJob,
    Project,
    ReviewItem,
    RouteAttempt,
    SourceFile,
    Tenant,
    User,
    utcnow,
)
from akc_api.services import credit_entry, run_compile_job
from akc_api.settings import Settings
from akc_api.visual_gpu import validate_visual_result, visual_attestation
from akc_worker_document.settings import AnalysisWorkerSettings
from akc_worker_document.worker import (
    AnalysisAttemptError,
    AnalysisRuntime,
    AnalysisWorker,
    analysis_claim_statement,
)
from docx import Document as WordDocument
from PIL import Image
from pydantic import ValidationError
from pypdf import PdfWriter
from sqlalchemy import func, select
from sqlalchemy.dialects import postgresql

_TEST_SUPPORT_KEY = "analysis-isolation-verification-key"


@pytest_asyncio.fixture
async def analysis_api(
    tmp_path: Path,
) -> AsyncIterator[tuple[httpx.AsyncClient, Any, Settings]]:
    settings = Settings(
        env="test",
        database_url=f"sqlite+aiosqlite:///{(tmp_path / 'analysis.db').as_posix()}",
        data_dir=tmp_path / "data",
        local_background_tasks=False,
        local_analysis_worker_enabled=False,
        analysis_max_source_bytes=1024 * 1024,
        analysis_max_attempts=2,
        analysis_lease_seconds=40,
        analysis_attempt_timeout_seconds=30,
        analysis_backoff_base_seconds=0.01,
        analysis_backoff_max_seconds=0.02,
        clamav_enabled=False,
        allow_development_antivirus_bypass=True,
        collection_metadata_encryption_enabled=True,
        test_support_key=_TEST_SUPPORT_KEY,
    )
    app = create_app(settings)
    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app, raise_app_exceptions=True)
        async with httpx.AsyncClient(
            transport=transport,
            base_url="http://testserver",
        ) as client:
            yield client, app, settings


async def _register(
    client: httpx.AsyncClient,
    *,
    email: str = "analysis-owner@example.com",
    tenant_name: str = "Analysis Isolation",
) -> dict[str, Any]:
    response = await client.post(
        "/v1/auth/register",
        json={
            "email": email,
            "password": "correct horse battery staple",
            "display_name": "Analysis Owner",
            "tenant_name": tenant_name,
        },
    )
    assert response.status_code == 201, response.text
    assert response.json()["email_verified"] is False
    captured = await client.post(
        "/__test__/verification-token",
        headers={"X-AKC-Test-Support-Key": _TEST_SUPPORT_KEY},
        json={"email": email},
    )
    assert captured.status_code == 200
    verified = await client.post(
        "/v1/auth/verify-email",
        json={"token": captured.json()["token"]},
    )
    assert verified.status_code == 200, verified.text
    return verified.json()


async def _create_project(client: httpx.AsyncClient) -> str:
    response = await client.post(
        "/v1/projects",
        json={"name": f"Sandbox {uuid.uuid4()}", "description": "Isolation evidence"},
    )
    assert response.status_code == 201, response.text
    return str(response.json()["id"])


async def _upload(
    client: httpx.AsyncClient,
    *,
    project_id: str,
    filename: str,
    content_type: str,
    content: bytes,
) -> str:
    digest = hashlib.sha256(content).hexdigest()
    initiated = await client.post(
        "/v1/uploads/initiate",
        json={
            "project_id": project_id,
            "filename": filename,
            "size": len(content),
            "content_type": content_type,
            "sha256": digest,
        },
    )
    assert initiated.status_code == 201, initiated.text
    target = initiated.json()
    uploaded = await client.put(
        target["upload_url"],
        content=content,
        headers=target["headers"],
    )
    assert uploaded.status_code == 204, uploaded.text
    completed = await client.post(
        f"/v1/uploads/{target['upload_id']}/complete",
        json={"sha256": digest},
    )
    assert completed.status_code == 200, completed.text
    return str(target["document_id"])


async def _enqueue(client: httpx.AsyncClient, document_id: str) -> uuid.UUID:
    response = await client.post(f"/v1/documents/{document_id}/analyze")
    assert response.status_code == 202, response.text
    assert response.json()["status"] == "queued"
    return uuid.UUID(response.json()["task_id"])


def _worker(app: Any, settings: Settings) -> AnalysisWorker:
    return AnalysisWorker(
        engine=app.state.database.engine,
        store=app.state.object_store,
        runtime=AnalysisRuntime.from_api_settings(settings),
    )


async def test_enqueue_is_naturally_idempotent_and_rejects_oversize_before_upload(
    analysis_api: tuple[httpx.AsyncClient, Any, Settings],
) -> None:
    client, app, settings = analysis_api
    registration = await _register(client)
    tenant_id = uuid.UUID(registration["tenant_id"])
    project_id = await _create_project(client)

    async with app.state.database.sessions.begin() as session:
        tenant = await session.get(Tenant, tenant_id)
        assert tenant is not None
        tenant.plan_code = "team"
    rejected = await client.post(
        "/v1/uploads/initiate",
        json={
            "project_id": project_id,
            "filename": "too-large.pdf",
            "size": settings.analysis_max_source_bytes + 1,
            "content_type": "application/pdf",
            "sha256": "0" * 64,
        },
    )
    assert rejected.status_code == 413
    assert rejected.json()["error"]["code"] == "ANALYSIS_SOURCE_TOO_LARGE"
    assert rejected.json()["error"]["details"]["max_bytes"] == settings.analysis_max_source_bytes

    document_id = await _upload(
        client,
        project_id=project_id,
        filename="idempotent.txt",
        content_type="text/plain",
        content=b"one durable task and one outbox event",
    )
    first = await client.post(f"/v1/documents/{document_id}/analyze")
    second = await client.post(f"/v1/documents/{document_id}/analyze")
    assert first.status_code == second.status_code == 202
    assert first.json()["task_id"] == second.json()["task_id"]
    assert first.json()["status"] == second.json()["status"] == "queued"

    task_id = uuid.UUID(first.json()["task_id"])
    async with app.state.database.sessions() as session:
        task_count = await session.scalar(
            select(func.count(AnalysisTask.id)).where(AnalysisTask.id == task_id)
        )
        outbox_count = await session.scalar(
            select(func.count(OutboxEvent.id)).where(
                OutboxEvent.aggregate_id == task_id,
                OutboxEvent.event_type == "document.analysis.requested.v1",
            )
        )
    assert task_count == 1
    assert outbox_count == 1


async def test_single_worker_concurrency_persists_one_result_and_tenant_safe_preview(
    analysis_api: tuple[httpx.AsyncClient, Any, Settings],
) -> None:
    client, app, settings = analysis_api
    await _register(client)
    project_id = await _create_project(client)
    document_id = await _upload(
        client,
        project_id=project_id,
        filename="preview.txt",
        content_type="text/plain",
        content=b"Rendered preview with a provenance-aligned bounding box.",
    )
    task_id = await _enqueue(client, document_id)
    worker = _worker(app, settings)

    outcomes = await asyncio.gather(
        worker.run_once(task_id=task_id),
        worker.run_once(task_id=task_id),
    )
    assert sorted(outcomes) == [False, True]
    assert await worker.run_once(task_id=task_id) is False

    async with app.state.database.sessions() as session:
        task = await session.get(AnalysisTask, task_id)
        pages = list(
            await session.scalars(select(Page).where(Page.document_id == uuid.UUID(document_id)))
        )
        blocks = list(
            await session.scalars(select(Block).where(Block.document_id == uuid.UUID(document_id)))
        )
        assets = (
            list(await session.scalars(select(PageAsset).where(PageAsset.page_id == pages[0].id)))
            if pages
            else []
        )
    assert task is not None
    assert task.status == "completed"
    assert task.attempt_count == 1
    assert len(pages) == len(blocks) == 1
    assert blocks[0].bbox1000 == [1, 1, 999, 999]
    assert {asset.asset_type for asset in assets} == {
        "preview",
        "thumbnail",
        "inference_raster",
    }
    inference = [asset for asset in assets if asset.asset_type == "inference_raster"]
    assert {asset.metadata_json["dpi"] for asset in inference} == {200, 300}
    assert all(
        asset.metadata_json["colorspace"] == "RGB" and asset.metadata_json["page_index0"] == 0
        for asset in inference
    )

    preview = await client.get(f"/v1/pages/{pages[0].id}/preview")
    assert preview.status_code == 200
    assert preview.content.startswith(b"\x89PNG\r\n\x1a\n")
    assert preview.headers["cache-control"].startswith("private")
    assert preview.headers["etag"].startswith('"sha256-')

    await _register(
        client,
        email="other-tenant@example.com",
        tenant_name="Other Preview Tenant",
    )
    cross_tenant = await client.get(f"/v1/pages/{pages[0].id}/preview")
    assert cross_tenant.status_code == 404
    login = await client.post(
        "/v1/auth/login",
        json={
            "email": "analysis-owner@example.com",
            "password": "correct horse battery staple",
        },
    )
    assert login.status_code == 200
    assert (await client.get(f"/v1/pages/{pages[0].id}/preview")).status_code == 200


async def test_shared_task_pause_is_binding_scoped_and_reuse_is_never_double_billed(
    analysis_api: tuple[httpx.AsyncClient, Any, Settings],
) -> None:
    client, app, settings = analysis_api
    registration = await _register(
        client,
        email="shared-collection-owner@example.com",
        tenant_name="Shared Collection Runtime",
    )
    tenant_id = uuid.UUID(registration["tenant_id"])
    project_id = uuid.UUID(await _create_project(client))
    document_id = uuid.UUID(
        await _upload(
            client,
            project_id=str(project_id),
            filename="shared-result.txt",
            content_type="text/plain",
            content=b"One extraction result reused by two collection jobs.",
        )
    )
    task_id = await _enqueue(client, str(document_id))
    resume_token = uuid.uuid4().hex + uuid.uuid4().hex
    binding_ids: list[uuid.UUID] = []
    collection_ids: list[uuid.UUID] = []
    job_ids: list[uuid.UUID] = []

    async with app.state.database.sessions.begin() as session:
        user = await session.get(User, uuid.UUID(registration["user_id"]))
        task = await session.get(AnalysisTask, task_id)
        source = await session.get(SourceFile, task.source_file_id if task else uuid.uuid4())
        assert user is not None and task is not None and source is not None
        for index in range(2):
            collection_id = uuid.uuid4()
            root_id = uuid.uuid4()
            collection_file_id = uuid.uuid4()
            job_id = uuid.uuid4()
            plan_id = uuid.uuid4()
            collection = Collection(
                id=collection_id,
                tenant_id=tenant_id,
                project_id=project_id,
                name=f"Shared collection {index}",
                status="PROCESSING" if index == 0 else "PAUSED",
                paused_from=None if index == 0 else "PROCESSING",
                profile={},
                manifest_revision=1,
                created_by=user.id,
            )
            root_name = app.state.collection_metadata_codec.encrypt_source_root_display_name(
                f"Shared root {index}",
                tenant_id=tenant_id,
                collection_id=collection_id,
                source_root_id=root_id,
            )
            root = CollectionSourceRoot(
                id=root_id,
                tenant_id=tenant_id,
                collection_id=collection_id,
                display_name_ciphertext=root_name.ciphertext,
                metadata_key_id=root_name.key_id,
                source_fingerprint=hashlib.sha256(f"root:{index}".encode()).hexdigest(),
                created_by=user.id,
            )
            relative_path = f"shared/{index}.txt"
            file_path = app.state.collection_metadata_codec.encrypt_file_relative_path(
                relative_path,
                tenant_id=tenant_id,
                collection_id=collection_id,
                source_root_id=root_id,
                file_id=collection_file_id,
            )
            file_name = app.state.collection_metadata_codec.encrypt_file_display_name(
                f"{index}.txt",
                tenant_id=tenant_id,
                collection_id=collection_id,
                source_root_id=root_id,
                file_id=collection_file_id,
            )
            blind_index = app.state.collection_metadata_codec.relative_path_blind_index(
                relative_path,
                tenant_id=tenant_id,
                collection_id=collection_id,
                source_root_id=root_id,
            )
            collection_file = CollectionFile(
                id=collection_file_id,
                tenant_id=tenant_id,
                collection_id=collection_id,
                source_root_id=root_id,
                source_file_id=source.id,
                relative_path_ciphertext=file_path.ciphertext,
                display_name_ciphertext=file_name.ciphertext,
                metadata_key_id=file_path.key_id,
                relative_path_blind_index=blind_index.digest,
                relative_path_blind_index_key_id=blind_index.key_id,
                size_bytes=source.size_bytes,
                expected_mime="text/plain",
                detected_mime="text/plain",
                sha256=source.sha256,
                status="verified",
            )
            job = ProcessingJob(
                id=job_id,
                tenant_id=tenant_id,
                project_id=project_id,
                document_id=None,
                job_type="collection_processing",
                status="running" if index == 0 else "paused",
                requested_options={
                    "architecture_plan_id": str(plan_id),
                    "immutable_plan_sha256": "a" * 64,
                    "approved_preflight_sha256": "b" * 64,
                    "approved_estimate_sha256": "c" * 64,
                    "resume_token_hash": hashlib.sha256(resume_token.encode()).hexdigest(),
                    "resume_version": 1,
                },
                progress={
                    "stage": "processing" if index == 0 else "paused",
                    "total_tasks": 1,
                    "completed_tasks": 0,
                    "failed_tasks": 0,
                    "terminal_result_ids": [],
                },
                cost_estimate={
                    "hard_cap": "0.000000",
                    "overage_policy": "stop_at_cap",
                },
                cost_actual={
                    "reserved": "0.000000",
                    "consumed": "0.000000",
                    "refunded": "0.000000",
                    "released": "0.000000",
                    "billable_pages": 0,
                    "unbillable_pages": 0,
                },
                started_at=utcnow(),
            )
            plan = ArchitecturePlan(
                id=plan_id,
                tenant_id=tenant_id,
                collection_id=collection_id,
                processing_job_id=job_id,
                plan_version=1,
                status="planned",
                input_integrity_sha256="a" * 64,
                plan={"execution_scope": "collection_processing_runtime"},
                created_by=user.id,
            )
            binding = CollectionProcessingTaskBinding(
                tenant_id=tenant_id,
                collection_id=collection_id,
                processing_job_id=job_id,
                analysis_task_id=task.id,
                collection_file_id=collection_file_id,
                document_id=document_id,
                billing_disposition="reuse_unbillable",
                billing_owner_job_id=None,
                billing_basis_sha256=hashlib.sha256(b"shared-result").hexdigest(),
                status="active" if index == 0 else "paused",
            )
            session.add(collection)
            await session.flush()
            session.add(root)
            await session.flush()
            session.add_all((collection_file, job))
            await session.flush()
            session.add_all((plan, binding))
            await session.flush()
            binding_ids.append(binding.id)
            collection_ids.append(collection_id)
            job_ids.append(job_id)

    worker = _worker(app, settings)
    assert await worker.run_once(task_id=task_id) is True

    async with app.state.database.sessions() as session:
        first_binding = await session.get(CollectionProcessingTaskBinding, binding_ids[0])
        paused_binding = await session.get(CollectionProcessingTaskBinding, binding_ids[1])
        first_collection = await session.get(Collection, collection_ids[0])
        paused_collection = await session.get(Collection, collection_ids[1])
        first_routes = list(
            await session.scalars(
                select(RouteAttempt).where(RouteAttempt.collection_id == collection_ids[0])
            )
        )
        paused_routes = list(
            await session.scalars(
                select(RouteAttempt).where(RouteAttempt.collection_id == collection_ids[1])
            )
        )
    assert first_binding is not None and first_binding.status == "settled"
    assert paused_binding is not None and paused_binding.status == "paused"
    assert first_collection is not None and first_collection.status == "VERIFYING_OUTPUT"
    assert paused_collection is not None and paused_collection.status == "PAUSED"
    assert first_routes and all(route.actual_credits == 0 for route in first_routes)
    assert paused_routes == []

    dashboard = await client.get("/v1/dashboard")
    assert dashboard.status_code == 200, dashboard.text
    assert dashboard.json()["active_jobs"] == 2

    resumed = await client.post(
        f"/v1/collections/{collection_ids[1]}/processing/control",
        headers={"Idempotency-Key": "resume-shared-binding"},
        json={"action": "resume", "processing_resume_token": resume_token},
    )
    assert resumed.status_code == 200, resumed.text
    assert resumed.json()["processing_job_id"] == str(job_ids[1])
    assert resumed.json()["credits_consumed"] == "0.000000"

    async with app.state.database.sessions() as session:
        resumed_binding = await session.get(CollectionProcessingTaskBinding, binding_ids[1])
        resumed_collection = await session.get(Collection, collection_ids[1])
        resumed_routes = list(
            await session.scalars(
                select(RouteAttempt).where(RouteAttempt.collection_id == collection_ids[1])
            )
        )
        consumes = int(
            await session.scalar(
                select(func.count(CreditLedger.id)).where(
                    CreditLedger.job_id.in_(job_ids),
                    CreditLedger.entry_type == "consume",
                )
            )
            or 0
        )
    assert resumed_binding is not None and resumed_binding.status == "settled"
    assert resumed_collection is not None
    assert resumed_collection.status == "VERIFYING_OUTPUT"
    assert resumed_routes and all(route.actual_credits == 0 for route in resumed_routes)
    assert consumes == 0


async def test_credit_failure_retry_versions_plan_and_reconciles_ledger_once(
    analysis_api: tuple[httpx.AsyncClient, Any, Settings],
) -> None:
    client, app, settings = analysis_api
    registration = await _register(
        client,
        email="collection-retry-owner@example.com",
        tenant_name="Collection Retry Runtime",
    )
    tenant_id = uuid.UUID(registration["tenant_id"])
    user_id = uuid.UUID(registration["user_id"])
    project_id = uuid.UUID(await _create_project(client))
    document_id = uuid.UUID(
        await _upload(
            client,
            project_id=str(project_id),
            filename="retry-result.txt",
            content_type="text/plain",
            content=b"A completed extraction result that exceeded its collection hard cap.",
        )
    )
    task_id = await _enqueue(client, str(document_id))
    worker = _worker(app, settings)
    assert await worker.run_once(task_id=task_id) is True

    collection_id = uuid.uuid4()
    collection_file_id = uuid.uuid4()
    plan_id = uuid.uuid4()
    job_id = uuid.uuid4()
    module_keys = (
        "source_index",
        "document_catalog",
        "knowledge_notes",
        "entities",
        "relations",
        "integrity",
        "export_manifest",
    )
    async with app.state.database.sessions.begin() as session:
        task = await session.get(AnalysisTask, task_id)
        source = await session.get(SourceFile, task.source_file_id if task else uuid.uuid4())
        page = await session.scalar(
            select(Page).where(Page.tenant_id == tenant_id, Page.document_id == document_id)
        )
        assert task is not None and source is not None and page is not None
        page.status = "COMPLETED"
        page.route = "native"
        page.preflight_metrics = {
            **(page.preflight_metrics or {}),
            "actual_credits": "0.750000",
        }
        root_id = uuid.uuid4()
        root_name = app.state.collection_metadata_codec.encrypt_source_root_display_name(
            "Retry root",
            tenant_id=tenant_id,
            collection_id=collection_id,
            source_root_id=root_id,
        )
        root = CollectionSourceRoot(
            id=root_id,
            tenant_id=tenant_id,
            collection_id=collection_id,
            display_name_ciphertext=root_name.ciphertext,
            metadata_key_id=root_name.key_id,
            source_fingerprint=hashlib.sha256(b"retry-root").hexdigest(),
            created_by=user_id,
        )
        collection = Collection(
            id=collection_id,
            tenant_id=tenant_id,
            project_id=project_id,
            name="Retry collection",
            status="FAILED_RETRYABLE",
            status_reason="CREDIT_HARD_CAP_REACHED",
            profile={},
            manifest_revision=1,
            created_by=user_id,
        )
        session.add(collection)
        await session.flush()
        session.add(root)
        await session.flush()
        relative_path = "retry/result.txt"
        file_path = app.state.collection_metadata_codec.encrypt_file_relative_path(
            relative_path,
            tenant_id=tenant_id,
            collection_id=collection_id,
            source_root_id=root.id,
            file_id=collection_file_id,
        )
        file_name = app.state.collection_metadata_codec.encrypt_file_display_name(
            "result.txt",
            tenant_id=tenant_id,
            collection_id=collection_id,
            source_root_id=root.id,
            file_id=collection_file_id,
        )
        blind_index = app.state.collection_metadata_codec.relative_path_blind_index(
            relative_path,
            tenant_id=tenant_id,
            collection_id=collection_id,
            source_root_id=root.id,
        )
        collection_file = CollectionFile(
            id=collection_file_id,
            tenant_id=tenant_id,
            collection_id=collection_id,
            source_root_id=root.id,
            source_file_id=source.id,
            relative_path_ciphertext=file_path.ciphertext,
            display_name_ciphertext=file_name.ciphertext,
            metadata_key_id=file_path.key_id,
            relative_path_blind_index=blind_index.digest,
            relative_path_blind_index_key_id=blind_index.key_id,
            size_bytes=source.size_bytes,
            expected_mime="text/plain",
            detected_mime="text/plain",
            sha256=source.sha256,
            status="verified",
        )
        module_ids = {key: uuid.uuid4() for key in module_keys}
        output_modules = [
            {"id": str(module_ids[key]), "module_key": key, "module_version": "1.0"}
            for key in module_keys
        ]
        plan_payload = {
            "schema_version": "1.0",
            "execution_scope": "collection_processing_runtime",
            "collection_id": str(collection_id),
            "project_id": str(project_id),
            "processing_job_id": str(job_id),
            "credit_hard_cap": "0.500000",
            "reserve_ceiling": "0.500000",
            "approved_collection_reserve_ceiling": "1.000000",
            "overage_policy": "stop_at_cap",
            "knowledge_blueprint_id": "generic-mixed-corpus",
            "knowledge_blueprint_registry_sha256": "sha256:" + "d" * 64,
            "knowledge_blueprint_module_sha256": "sha256:" + "e" * 64,
            "approved_preflight_sha256": "b" * 64,
            "approved_estimate_sha256": "c" * 64,
            "output_modules": output_modules,
            "analysis_task_ids": [str(task.id)],
        }
        plan_sha = hashlib.sha256(
            (json.dumps(plan_payload, separators=(",", ":"), sort_keys=True) + "\n").encode()
        ).hexdigest()
        plan_payload["immutable_plan_sha256"] = plan_sha
        job = ProcessingJob(
            id=job_id,
            tenant_id=tenant_id,
            project_id=project_id,
            document_id=None,
            job_type="collection_processing",
            status="failed",
            requested_options={
                **plan_payload,
                "architecture_plan_id": str(plan_id),
                "initiating_user_id": str(user_id),
                "resume_token_hash": hashlib.sha256(b"superseded-resume-token").hexdigest(),
                "resume_version": 1,
            },
            progress={
                "stage": "failed",
                "total_tasks": 1,
                "completed_tasks": 1,
                "failed_tasks": 0,
                "terminal_result_ids": [str(page.id)],
            },
            cost_estimate={
                "reserve_ceiling": "0.500000",
                "approved_collection_reserve_ceiling": "1.000000",
                "hard_cap": "0.500000",
                "overage_policy": "stop_at_cap",
            },
            cost_actual={
                "reserved": "0.500000",
                "consumed": "0.000000",
                "refunded": "0.000000",
                "released": "0.500000",
                "billable_pages": 0,
                "unbillable_pages": 1,
            },
            error={"code": "CREDIT_HARD_CAP_REACHED"},
            started_at=utcnow(),
            completed_at=utcnow(),
        )
        session.add_all((collection_file, job))
        await session.flush()
        plan = ArchitecturePlan(
            id=plan_id,
            tenant_id=tenant_id,
            collection_id=collection_id,
            processing_job_id=job_id,
            plan_version=1,
            status="planned",
            input_integrity_sha256=plan_sha,
            plan=plan_payload,
            created_by=user_id,
        )
        binding = CollectionProcessingTaskBinding(
            tenant_id=tenant_id,
            collection_id=collection_id,
            processing_job_id=job_id,
            analysis_task_id=task.id,
            collection_file_id=collection_file_id,
            document_id=document_id,
            billing_disposition="new_billable",
            billing_owner_job_id=job_id,
            billing_basis_sha256=hashlib.sha256(b"retry-billing-basis").hexdigest(),
            status="settled",
            settled_at=utcnow(),
        )
        session.add(plan)
        await session.flush()
        session.add_all(
            [
                BlueprintModule(
                    id=module_ids[key],
                    tenant_id=tenant_id,
                    collection_id=collection_id,
                    architecture_plan_id=plan.id,
                    module_key=key,
                    module_version="1.0",
                    status="planned",
                    config_json={},
                    output_summary={},
                )
                for key in module_keys
            ]
            + [
                binding,
                RouteAttempt(
                    tenant_id=tenant_id,
                    collection_id=collection_id,
                    collection_file_id=collection_file_id,
                    page_id=page.id,
                    attempt_number=1,
                    route=str(page.route or "native"),
                    status="unresolved",
                    reason_codes=["CREDIT_HARD_CAP_REACHED"],
                    estimated_credits=Decimal("0.750000"),
                    actual_credits=Decimal("0.000000"),
                    completed_at=utcnow(),
                ),
            ]
        )
        await credit_entry(
            session,
            tenant_id=tenant_id,
            operation_key=f"collection:{job_id}:reserve",
            entry_type="reserve",
            credits=Decimal("0.500000"),
            job_id=job_id,
        )
        await credit_entry(
            session,
            tenant_id=tenant_id,
            operation_key=f"collection:{job_id}:release",
            entry_type="release",
            credits=Decimal("0.500000"),
            job_id=job_id,
        )

    retry_headers = {"Idempotency-Key": "retry-credit-terminal"}
    retried = await client.post(
        f"/v1/collections/{collection_id}/processing/retry",
        headers=retry_headers,
        json={"credit_hard_cap": "1.000000"},
    )
    replayed = await client.post(
        f"/v1/collections/{collection_id}/processing/retry",
        headers=retry_headers,
        json={"credit_hard_cap": "1.000000"},
    )
    assert retried.status_code == replayed.status_code == 200, retried.text
    assert retried.json() == replayed.json()
    body = retried.json()
    assert body["collection_status"] == "VERIFYING_OUTPUT"
    assert body["processing_status"] == "running"
    assert body["credit_hard_cap"] == "1.000000"
    assert body["processing_resume_token"] is not None

    async with app.state.database.sessions() as session:
        plans = list(
            await session.scalars(
                select(ArchitecturePlan)
                .where(ArchitecturePlan.collection_id == collection_id)
                .order_by(ArchitecturePlan.plan_version)
            )
        )
        binding = await session.scalar(
            select(CollectionProcessingTaskBinding).where(
                CollectionProcessingTaskBinding.processing_job_id == job_id
            )
        )
        route_attempts = list(
            await session.scalars(
                select(RouteAttempt).where(RouteAttempt.collection_id == collection_id)
            )
        )
        ledgers = list(
            await session.scalars(
                select(CreditLedger)
                .where(CreditLedger.job_id == job_id)
                .order_by(CreditLedger.created_at, CreditLedger.id)
            )
        )
        finalizer_events = list(
            await session.scalars(
                select(OutboxEvent).where(
                    OutboxEvent.aggregate_id == job_id,
                    OutboxEvent.event_type == "collection.semantic.compile.requested.v1",
                )
            )
        )
    assert [plan.status for plan in plans] == ["stale", "planned"]
    assert plans[1].plan["retry_of_architecture_plan_id"] == str(plans[0].id)
    assert plans[1].plan["credit_hard_cap"] == "1.000000"
    assert binding is not None and binding.status == "settled"
    assert len(route_attempts) == 1
    assert route_attempts[0].status == "verified"
    assert route_attempts[0].actual_credits == Decimal("0.750000")
    assert [ledger.entry_type for ledger in ledgers].count("consume") == 1
    assert sum(
        (ledger.credits for ledger in ledgers if ledger.entry_type == "consume"),
        Decimal("0"),
    ) == Decimal("0.750000")
    assert len(finalizer_events) == 1
    assert finalizer_events[0].payload["architecture_plan_id"] == str(plans[1].id)


async def test_package_failure_retry_reconsumes_refund_and_redrives_once(
    analysis_api: tuple[httpx.AsyncClient, Any, Settings],
) -> None:
    client, app, _settings = analysis_api
    registration = await _register(
        client,
        email="package-retry-owner@example.com",
        tenant_name="Package Retry Runtime",
    )
    tenant_id = uuid.UUID(registration["tenant_id"])
    user_id = uuid.UUID(registration["user_id"])
    project_id = uuid.UUID(await _create_project(client))
    collection_id = uuid.uuid4()
    job_id = uuid.uuid4()
    plan_id = uuid.uuid4()
    async with app.state.database.sessions.begin() as session:
        collection = Collection(
            id=collection_id,
            tenant_id=tenant_id,
            project_id=project_id,
            name="Package retry collection",
            status="FAILED_RETRYABLE",
            status_reason="COLLECTION_PACKAGE_STORAGE_FAILED",
            profile={},
            manifest_revision=1,
            created_by=user_id,
        )
        job = ProcessingJob(
            id=job_id,
            tenant_id=tenant_id,
            project_id=project_id,
            document_id=None,
            job_type="collection_processing",
            status="failed",
            requested_options={
                "execution_scope": "collection_processing_runtime",
                "architecture_plan_id": str(plan_id),
                "immutable_plan_sha256": "a" * 64,
                "approved_preflight_sha256": "b" * 64,
                "approved_estimate_sha256": "c" * 64,
                "initiating_user_id": str(user_id),
            },
            progress={
                "stage": "package_failed",
                "package_attempt": 0,
                "total_tasks": 0,
                "completed_tasks": 0,
                "failed_tasks": 0,
            },
            cost_estimate={
                "reserve_ceiling": "1.000000",
                "hard_cap": "1.000000",
                "overage_policy": "stop_at_cap",
            },
            cost_actual={
                "reserved": "1.500000",
                "consumed": "1.000000",
                "refunded": "1.000000",
                "released": "0.000000",
                "billable_pages": 1,
                "unbillable_pages": 0,
            },
            error={"code": "COLLECTION_PACKAGE_STORAGE_FAILED"},
            started_at=utcnow(),
            completed_at=utcnow(),
        )
        session.add_all((collection, job))
        await session.flush()
        session.add(
            ArchitecturePlan(
                id=plan_id,
                tenant_id=tenant_id,
                collection_id=collection_id,
                processing_job_id=job_id,
                plan_version=1,
                status="compiled",
                input_integrity_sha256="a" * 64,
                plan={"execution_scope": "collection_processing_runtime"},
                created_by=user_id,
            )
        )
        await credit_entry(
            session,
            tenant_id=tenant_id,
            operation_key=f"collection:{job_id}:reserve",
            entry_type="reserve",
            credits=Decimal("1.500000"),
            job_id=job_id,
        )
        await credit_entry(
            session,
            tenant_id=tenant_id,
            operation_key=f"collection:{job_id}:page:original:consume",
            entry_type="consume",
            credits=Decimal("1.000000"),
            job_id=job_id,
            metadata={"from_reserved": True},
        )
        await credit_entry(
            session,
            tenant_id=tenant_id,
            operation_key=f"collection:{job_id}:package-refund",
            entry_type="refund",
            credits=Decimal("1.000000"),
            job_id=job_id,
        )

    headers = {"Idempotency-Key": "retry-package-terminal"}
    retried = await client.post(
        f"/v1/collections/{collection_id}/processing/retry",
        headers=headers,
        json={},
    )
    replayed = await client.post(
        f"/v1/collections/{collection_id}/processing/retry",
        headers=headers,
        json={},
    )
    assert retried.status_code == replayed.status_code == 200, retried.text
    assert retried.json() == replayed.json()
    assert retried.json()["collection_status"] == "PACKAGING"
    assert retried.json()["processing_status"] == "running"
    assert retried.json()["credits_consumed"] == "2.000000"
    assert retried.json()["credits_refunded"] == "1.000000"

    async with app.state.database.sessions() as session:
        job = await session.get(ProcessingJob, job_id)
        ledgers = list(
            await session.scalars(
                select(CreditLedger)
                .where(CreditLedger.job_id == job_id)
                .order_by(CreditLedger.created_at, CreditLedger.id)
            )
        )
        finalizer_events = list(
            await session.scalars(
                select(OutboxEvent).where(
                    OutboxEvent.aggregate_id == job_id,
                    OutboxEvent.event_type == "collection.semantic.compile.requested.v1",
                )
            )
        )
    assert job is not None and job.progress["stage"] == "ready_for_packaging"
    assert job.progress["package_attempt"] == 1
    assert job.cost_actual["retry_reconsumed"] == "1.000000"
    assert sum(ledger.entry_type == "consume" for ledger in ledgers) == 2
    assert sum(ledger.entry_type == "refund" for ledger in ledgers) == 1
    assert len(finalizer_events) == 1

    deleted = await client.delete(
        f"/v1/collections/{collection_id}",
        headers={"Idempotency-Key": "delete-package-retry-before-finalizer"},
    )
    assert deleted.status_code == 200, deleted.text
    async with app.state.database.sessions() as session:
        cancelled_job = await session.get(ProcessingJob, job_id)
        cancelled_event = await session.get(OutboxEvent, finalizer_events[0].id)
        cancellation_releases = list(
            await session.scalars(
                select(CreditLedger).where(
                    CreditLedger.job_id == job_id,
                    CreditLedger.entry_type == "release",
                )
            )
        )
    assert cancelled_job is not None and cancelled_job.status == "cancelled"
    assert cancelled_job.progress["cancelled_finalizer_events"] == 1
    assert cancelled_job.cost_actual["released"] == "0.500000"
    assert cancelled_event is not None and cancelled_event.dead_lettered_at is not None
    assert len(cancellation_releases) == 1
    assert cancellation_releases[0].credits == Decimal("0.500000")


async def test_semantic_finalizer_dead_letter_is_customer_redrivable_once(
    analysis_api: tuple[httpx.AsyncClient, Any, Settings],
) -> None:
    client, app, _settings = analysis_api
    registration = await _register(
        client,
        email="finalizer-retry-owner@example.com",
        tenant_name="Finalizer Retry Runtime",
    )
    tenant_id = uuid.UUID(registration["tenant_id"])
    user_id = uuid.UUID(registration["user_id"])
    project_id = uuid.UUID(await _create_project(client))
    collection_id = uuid.uuid4()
    job_id = uuid.uuid4()
    plan_id = uuid.uuid4()
    async with app.state.database.sessions.begin() as session:
        collection = Collection(
            id=collection_id,
            tenant_id=tenant_id,
            project_id=project_id,
            name="Finalizer retry collection",
            status="UNRESOLVED",
            status_reason="COLLECTION_FINALIZER_ATTEMPTS_EXHAUSTED",
            profile={},
            manifest_revision=1,
            created_by=user_id,
        )
        job = ProcessingJob(
            id=job_id,
            tenant_id=tenant_id,
            project_id=project_id,
            document_id=None,
            job_type="collection_processing",
            status="failed",
            requested_options={
                "execution_scope": "collection_processing_runtime",
                "architecture_plan_id": str(plan_id),
                "immutable_plan_sha256": "a" * 64,
                "approved_preflight_sha256": "b" * 64,
                "approved_estimate_sha256": "c" * 64,
                "initiating_user_id": str(user_id),
            },
            progress={
                "stage": "semantic_finalizer_failed",
                "total_tasks": 0,
                "completed_tasks": 0,
                "failed_tasks": 0,
            },
            cost_estimate={
                "reserve_ceiling": "1.000000",
                "hard_cap": "1.000000",
                "overage_policy": "stop_at_cap",
            },
            cost_actual={
                "reserved": "1.000000",
                "consumed": "1.000000",
                "refunded": "1.000000",
                "released": "0.000000",
                "billable_pages": 1,
                "unbillable_pages": 0,
            },
            error={
                "code": "COLLECTION_FINALIZER_ATTEMPTS_EXHAUSTED",
                "retryable": True,
            },
            started_at=utcnow(),
            completed_at=utcnow(),
        )
        session.add_all((collection, job))
        await session.flush()
        session.add(
            ArchitecturePlan(
                id=plan_id,
                tenant_id=tenant_id,
                collection_id=collection_id,
                processing_job_id=job_id,
                plan_version=1,
                status="planned",
                input_integrity_sha256="a" * 64,
                plan={"execution_scope": "collection_processing_runtime"},
                created_by=user_id,
            )
        )
        await credit_entry(
            session,
            tenant_id=tenant_id,
            operation_key=f"collection:{job_id}:reserve",
            entry_type="reserve",
            credits=Decimal("1.000000"),
            job_id=job_id,
        )
        await credit_entry(
            session,
            tenant_id=tenant_id,
            operation_key=f"collection:{job_id}:page:original:consume",
            entry_type="consume",
            credits=Decimal("1.000000"),
            job_id=job_id,
            metadata={"from_reserved": True},
        )
        await credit_entry(
            session,
            tenant_id=tenant_id,
            operation_key=f"collection:{job_id}:semantic-refund",
            entry_type="refund",
            credits=Decimal("1.000000"),
            job_id=job_id,
        )

    headers = {"Idempotency-Key": "retry-finalizer-dead-letter"}
    retried = await client.post(
        f"/v1/collections/{collection_id}/processing/retry",
        headers=headers,
        json={},
    )
    replayed = await client.post(
        f"/v1/collections/{collection_id}/processing/retry",
        headers=headers,
        json={},
    )
    assert retried.status_code == replayed.status_code == 200, retried.text
    assert retried.json() == replayed.json()
    assert retried.json()["collection_status"] == "VERIFYING_OUTPUT"
    assert retried.json()["processing_status"] == "running"

    async with app.state.database.sessions() as session:
        job = await session.get(ProcessingJob, job_id)
        events = list(
            await session.scalars(
                select(OutboxEvent).where(
                    OutboxEvent.aggregate_id == job_id,
                    OutboxEvent.event_type == "collection.semantic.compile.requested.v1",
                )
            )
        )
        ledgers = list(
            await session.scalars(select(CreditLedger).where(CreditLedger.job_id == job_id))
        )
    assert job is not None and job.progress["stage"] == "semantic_compile_queued"
    assert job.progress["finalizer_retry_attempt"] == 1
    assert job.cost_actual["retry_reconsumed"] == "1.000000"
    assert len(events) == 1
    assert events[0].payload["actor_user_id"] == str(user_id)
    assert sum(ledger.entry_type == "consume" for ledger in ledgers) == 2
    assert sum(ledger.entry_type == "refund" for ledger in ledgers) == 1


def _scanned_fixture(kind: str) -> tuple[str, str, bytes]:
    if kind == "image":
        payload = io.BytesIO()
        Image.new("RGB", (600, 800), "white").save(
            payload,
            format="PNG",
            dpi=(300, 300),
        )
        return "scan.png", "image/png", payload.getvalue()
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    payload = io.BytesIO()
    writer.write(payload)
    return "scan.pdf", "application/pdf", payload.getvalue()


@pytest.mark.parametrize("kind", ["image", "pdf"])
async def test_scanned_input_fake_provider_admits_exact_inference_page_once(
    analysis_api: tuple[httpx.AsyncClient, Any, Settings],
    kind: str,
) -> None:
    client, app, settings = analysis_api
    registration = await _register(
        client,
        email=f"visual-{kind}@example.com",
        tenant_name=f"Visual {kind}",
    )
    tenant_id = uuid.UUID(registration["tenant_id"])
    project_id = await _create_project(client)
    filename, content_type, content = _scanned_fixture(kind)
    document_id = await _upload(
        client,
        project_id=project_id,
        filename=filename,
        content_type=content_type,
        content=content,
    )
    async with app.state.database.sessions.begin() as session:
        tenant = await session.get(Tenant, tenant_id)
        project = await session.get(Project, uuid.UUID(project_id))
        assert tenant is not None and project is not None
        tenant.private_mode = False
        project.output_profile = {"processing_mode": "speed"}
        session.add_all(
            [
                FeatureFlag(
                    tenant_id=tenant_id,
                    key="paddle_fast_route",
                    enabled=True,
                    rollout_percent=100,
                ),
                ModelRegistry(
                    endpoint="paddle_fast_e2e",
                    model_id="paddle_fast_e2e_model",
                    revision="a" * 40,
                    runtime_image_digest="sha256:" + "b" * 64,
                    adapter_version="parser-adapter-1.0.0",
                    policy_version="router-scanned-e2e",
                    benchmark_report="benchmarks/paddle-fast-e2e.json",
                    enabled=True,
                    canary_percent=100,
                ),
            ]
        )
    task_id = await _enqueue(client, document_id)
    assert await _worker(app, settings).run_once(task_id=task_id) is True
    async with app.state.database.sessions() as session:
        page = await session.scalar(select(Page).where(Page.document_id == uuid.UUID(document_id)))
        task = await session.get(AnalysisTask, task_id)
        assert page is not None, (
            task.status if task is not None else None,
            task.last_error_code if task is not None else None,
        )
        original_page_status = page.status
        inference_assets = list(
            (
                await session.scalars(
                    select(PageAsset).where(
                        PageAsset.page_id == page.id,
                        PageAsset.asset_type == "inference_raster",
                    )
                )
            ).all()
        )
        assert {asset.metadata_json["dpi"] for asset in inference_assets} == {
            200,
            300,
        }

    compiled = await client.post(
        f"/v1/documents/{document_id}/compile",
        json={"route_profile": "parse_fast_v1"},
        headers={"Idempotency-Key": f"scanned-{kind}-fake-provider"},
    )
    assert compiled.status_code == 202, compiled.text
    job_id = uuid.UUID(compiled.json()["job_id"])
    async with app.state.database.sessions() as session:
        await run_compile_job(
            session=session,
            job_id=job_id,
            settings=app.state.settings,
            object_store=app.state.object_store,
        )
    async with app.state.database.sessions.begin() as session:
        job = await session.get(ProcessingJob, job_id)
        invocation = await session.scalar(
            select(GpuProviderInvocation).where(GpuProviderInvocation.job_id == job_id)
        )
        page = await session.scalar(select(Page).where(Page.document_id == uuid.UUID(document_id)))
        document = await session.get(Document, uuid.UUID(document_id))
        assert (
            job is not None and invocation is not None and page is not None and document is not None
        ), {
            "job_status": job.status if job is not None else None,
            "job_error": job.error if job is not None else None,
            "requested_options": (job.requested_options if job is not None else None),
            "page_route": page.route if page is not None else None,
            "page_status": page.status if page is not None else None,
        }
        page_asset = await session.get(
            PageAsset,
            uuid.UUID(invocation.options["page_asset_id"]),
        )
        assert page_asset is not None
        assert page_asset.asset_type == "inference_raster"
        assert invocation.input_object_key == page_asset.storage_key
        assert invocation.options["dpi"] == 200
        result_id = "sha256:" + "c" * 64
        output_payload = {
            "ok": True,
            "schema_version": "1.0",
            "result_id": result_id,
            "job_id": str(job.id),
            "tenant_id": str(tenant_id),
            "provider": invocation.provider_key,
            "worker_kind": "parser",
            "model_revision": invocation.model_revision,
            "runtime_image_digest": invocation.runtime_image_digest,
            "adapter_version": invocation.adapter_version,
            "input_sha256": f"sha256:{page_asset.sha256}",
            "input_bytes": page_asset.metadata_json["size_bytes"],
            "idempotency_key": invocation.idempotency_key,
            "idempotent_replay": False,
            "blocks": [
                {
                    "block_id": f"blk_scanned_{kind}",
                    "type": "paragraph",
                    "text": f"Verified scanned {kind} output",
                    "origin": "ocr_extracted",
                    "source_refs": [
                        {
                            "document_id": str(document.id),
                            "document_version_id": (f"{document.id}:v{document.active_version}"),
                            "page_index0": 0,
                            "page_number1": 1,
                            "bbox1000": [10, 20, 900, 200],
                        }
                    ],
                    "confidence": 0.99,
                    "token_confidences": [0.99, 0.98, 0.99],
                    "quality_flags": [],
                }
            ],
            "generated_claims": [],
            "warnings": [],
            "provider_metrics": {"block_count": 1},
            "provider_raw": {},
            "metrics": {"gpu_seconds": 1.0},
            "verification": None,
        }
        output_body = json.dumps(
            output_payload,
            separators=(",", ":"),
            sort_keys=True,
        ).encode()
        await app.state.object_store.put_derived(
            invocation.output_object_key,
            output_body,
        )
        result = validate_visual_result(
            output_payload=output_payload,
            expected_job_id=job.id,
            expected_tenant_id=tenant_id,
            expected_document_id=document.id,
            expected_document_version_id=invocation.document_version_id,
            expected_page_index0=0,
            expected_provider=invocation.provider_key,
            expected_model_revision=invocation.model_revision,
            expected_runtime_image_digest=invocation.runtime_image_digest,
            expected_adapter_version=invocation.adapter_version,
            expected_input_sha256=page_asset.sha256,
            expected_input_bytes=page_asset.metadata_json["size_bytes"],
            expected_idempotency_key=invocation.idempotency_key,
            expected_image_asset_id=page_asset.id,
        )
        manifest = {
            "schema_version": "1.0",
            "invocation_id": str(invocation.id),
            "job_id": str(job.id),
            "tenant_id": str(tenant_id),
            "provider": "runpod",
            "provider_job_id": f"fake-{kind}-provider",
            "endpoint_id": invocation.endpoint_id,
            "provider_key": invocation.provider_key,
            "model_revision": invocation.model_revision,
            "runtime_image_digest": invocation.runtime_image_digest,
            "adapter_version": invocation.adapter_version,
            "result_id": result_id,
            "output_object_key": invocation.output_object_key,
            "output_sha256": ("sha256:" + hashlib.sha256(output_body).hexdigest()),
            "output_bytes": len(output_body),
            "metrics": output_payload["metrics"],
            "warning_count": 0,
            "warning_sha256": [],
            "raw_provider_response_sha256": "sha256:" + "d" * 64,
            "completion_source": "poll",
            "visual_attestation": visual_attestation(
                result,
                page_id=page.id,
                page_index0=0,
                page_width_px=page_asset.metadata_json["width"],
                page_height_px=page_asset.metadata_json["height"],
                input_size_bytes=page_asset.metadata_json["size_bytes"],
            ),
        }
        invocation.status = "completed"
        invocation.provider_status = "COMPLETED"
        invocation.provider_job_id = f"fake-{kind}-provider"
        invocation.result_manifest = manifest
        invocation.result_manifest_sha256 = hashlib.sha256(
            json.dumps(
                manifest,
                separators=(",", ":"),
                sort_keys=True,
            ).encode()
        ).hexdigest()
        invocation.started_at = utcnow()
        invocation.completed_at = utcnow()

    async with app.state.database.sessions() as session:
        await run_compile_job(
            session=session,
            job_id=job_id,
            settings=app.state.settings,
            object_store=app.state.object_store,
        )
        await run_compile_job(
            session=session,
            job_id=job_id,
            settings=app.state.settings,
            object_store=app.state.object_store,
        )
    async with app.state.database.sessions() as session:
        job = await session.get(ProcessingJob, job_id)
        page = await session.scalar(select(Page).where(Page.document_id == uuid.UUID(document_id)))
        consumes = int(
            await session.scalar(
                select(func.count(CreditLedger.id)).where(
                    CreditLedger.job_id == job_id,
                    CreditLedger.entry_type == "consume",
                )
            )
            or 0
        )
        assert job is not None and job.status == "completed"
        assert page is not None and page.status == original_page_status
        assert consumes == 1


async def test_sensitive_preview_is_masked_by_default_and_opt_out_is_explicit(
    analysis_api: tuple[httpx.AsyncClient, Any, Settings],
) -> None:
    client, app, settings = analysis_api
    await _register(client)
    project_id = await _create_project(client)
    document_id = await _upload(
        client,
        project_id=project_id,
        filename="sensitive.txt",
        content_type="text/plain",
        content=(
            b"Contact the account owner at pii@example.com. "
            b"Rotate github_pat_12345678901234567890 immediately."
        ),
    )
    task_id = await _enqueue(client, document_id)
    assert await _worker(app, settings).run_once(task_id=task_id) is True

    async with app.state.database.sessions() as session:
        page = await session.scalar(select(Page).where(Page.document_id == uuid.UUID(document_id)))
        block = await session.scalar(
            select(Block).where(Block.document_id == uuid.UUID(document_id))
        )
        secret_review = await session.scalar(
            select(ReviewItem).where(
                ReviewItem.document_id == uuid.UUID(document_id),
                ReviewItem.category == "secret_detected",
            )
        )
    assert page is not None
    assert block is not None
    assert page.preflight_metrics["sensitive_data"] == {
        "has_pii": True,
        "has_secret": True,
        "external_transfer_requires_confirmation": True,
        "counts": {"api_key": 1, "email": 1},
        "detector_limitations": [
            "pattern detection can miss sensitive data",
            "pattern detection can produce false positives",
        ],
    }
    assert "mistral_fallback" not in page.preflight_metrics["ready_routes"]
    assert block.warnings == ["sensitive_api_key_detected", "sensitive_email_detected"]
    assert secret_review is None
    masked = await client.get(f"/v1/pages/{page.id}/preview")
    assert masked.status_code == 200
    assert masked.headers["x-akc-preview-redaction"] == "masked"
    assert masked.headers["x-akc-masked-regions"] == "1"
    with Image.open(io.BytesIO(masked.content)) as image:
        assert image.convert("RGB").getpixel((image.width // 2, image.height // 2)) == (
            0,
            0,
            0,
        )

    updated = await client.patch(
        "/v1/privacy",
        headers={"Idempotency-Key": str(uuid.uuid4())},
        json={"preview_pii_masking": False},
    )
    assert updated.status_code == 200, updated.text
    unmasked = await client.get(f"/v1/pages/{page.id}/preview")
    assert unmasked.status_code == 200
    assert unmasked.headers["x-akc-preview-redaction"] == "disabled"
    assert unmasked.headers["x-akc-masked-regions"] == "0"
    assert unmasked.headers["etag"] != masked.headers["etag"]
    with Image.open(io.BytesIO(unmasked.content)) as image:
        assert image.convert("RGB").getpixel((image.width // 2, image.height // 2)) == (
            255,
            255,
            255,
        )


async def test_hostile_text_is_dead_lettered_without_partial_rows(
    analysis_api: tuple[httpx.AsyncClient, Any, Settings],
) -> None:
    client, app, settings = analysis_api
    await _register(client)
    project_id = await _create_project(client)
    document_id = await _upload(
        client,
        project_id=project_id,
        filename="hostile.txt",
        content_type="text/plain",
        content=b"valid UTF-8 prefix\x00hidden binary suffix",
    )
    task_id = await _enqueue(client, document_id)
    assert await _worker(app, settings).run_once(task_id=task_id) is True

    async with app.state.database.sessions() as session:
        task = await session.get(AnalysisTask, task_id)
        document = await session.get(Document, uuid.UUID(document_id))
        event = await session.scalar(select(OutboxEvent).where(OutboxEvent.aggregate_id == task_id))
        page_count = await session.scalar(
            select(func.count(Page.id)).where(Page.document_id == uuid.UUID(document_id))
        )
    assert task is not None and task.status == "dead_letter"
    assert task.last_error_code == "BINARY_TEXT_FILE"
    assert task.attempt_count == 1
    assert document is not None and document.status == "PARSE_FAILED"
    assert event is not None
    assert event.last_error == "BINARY_TEXT_FILE"
    assert event.published_at is not None and event.dead_lettered_at is not None
    assert page_count == 0


async def test_timeout_retries_then_dead_letters_without_duplicate_pages(
    analysis_api: tuple[httpx.AsyncClient, Any, Settings],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    client, app, settings = analysis_api
    await _register(client)
    project_id = await _create_project(client)
    document_id = await _upload(
        client,
        project_id=project_id,
        filename="timeout.txt",
        content_type="text/plain",
        content=b"bounded retry evidence",
    )
    task_id = await _enqueue(client, document_id)
    worker = _worker(app, settings)

    async def timeout_sandbox(**_: object) -> object:
        raise AnalysisAttemptError("PARSER_TIMEOUT", retryable=True)

    monkeypatch.setattr(worker, "_invoke_sandbox", timeout_sandbox)
    assert await worker.run_once(task_id=task_id) is True
    async with app.state.database.sessions.begin() as session:
        task = await session.get(AnalysisTask, task_id)
        event = await session.scalar(select(OutboxEvent).where(OutboxEvent.aggregate_id == task_id))
        assert task is not None and task.status == "queued"
        assert task.attempt_count == 1
        assert task.last_error_code == "PARSER_TIMEOUT"
        assert event is not None and event.published_at is None
        task.available_at = utcnow()
        event.available_at = utcnow()

    assert await worker.run_once(task_id=task_id) is True
    assert await worker.run_once(task_id=task_id) is False
    async with app.state.database.sessions() as session:
        task = await session.get(AnalysisTask, task_id)
        event = await session.scalar(select(OutboxEvent).where(OutboxEvent.aggregate_id == task_id))
        pages = await session.scalar(
            select(func.count(Page.id)).where(Page.document_id == uuid.UUID(document_id))
        )
    assert task is not None and task.status == "dead_letter"
    assert task.attempt_count == task.max_attempts == 2
    assert task.last_error_code == "PARSER_TIMEOUT"
    assert event is not None
    assert event.published_at is not None and event.dead_lettered_at is not None
    assert pages == 0


async def test_office_analysis_is_explicitly_no_preview(
    analysis_api: tuple[httpx.AsyncClient, Any, Settings],
) -> None:
    client, app, settings = analysis_api
    await _register(client)
    project_id = await _create_project(client)
    buffer = io.BytesIO()
    word = WordDocument()
    word.add_heading("No synthetic Office preview", level=1)
    word.add_paragraph("The extracted content remains available with provenance.")
    word.save(buffer)
    document_id = await _upload(
        client,
        project_id=project_id,
        filename="office.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        content=buffer.getvalue(),
    )
    task_id = await _enqueue(client, document_id)
    assert await _worker(app, settings).run_once(task_id=task_id) is True

    async with app.state.database.sessions() as session:
        task = await session.get(AnalysisTask, task_id)
        page = await session.scalar(select(Page).where(Page.document_id == uuid.UUID(document_id)))
    assert task is not None and task.status == "completed"
    assert task.preview_count == 0
    assert page is not None
    assert page.thumbnail_key is None
    assert page.preflight_metrics["preview_unavailable_reason"] == "unsupported_document_preview"
    unavailable = await client.get(f"/v1/pages/{page.id}/preview")
    assert unavailable.status_code == 404
    assert unavailable.json()["error"]["details"]["reason"] == "unsupported_document_preview"


def test_production_import_graph_and_worker_resource_guards() -> None:
    statement = analysis_claim_statement(
        now=utcnow(),
        dialect_name="postgresql",
    )
    assert (
        "SKIP LOCKED"
        in str(
            statement.compile(
                dialect=postgresql.dialect(),
                compile_kwargs={"literal_binds": True},
            )
        ).upper()
    )

    completed = subprocess.run(
        [
            sys.executable,
            "-c",
            (
                "import sys; import akc_api.main; "
                "assert 'akc_api.parsers' not in sys.modules; "
                "assert 'akc_worker_document.sandbox_runner' not in sys.modules"
            ),
        ],
        check=False,
        capture_output=True,
        text=True,
        timeout=20,
    )
    assert completed.returncode == 0, completed.stderr

    with pytest.raises(ValidationError, match="in-process analysis adapter"):
        Settings(env="production", local_analysis_worker_enabled=True)
    with pytest.raises(ValidationError, match="child memory"):
        AnalysisWorkerSettings(
            analysis_child_memory_bytes=128 * 1024 * 1024,
        )
